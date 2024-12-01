# hr_data.py

import os
import pandas as pd
from pydantic import BaseModel, Field
from faker import Faker
import random
import uuid
import asyncio
import logging
from chatgpt import chat_with_instructor

fake = Faker()
Faker.seed(0)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='hr_data.log')

class Employee(BaseModel):
    employee_id: str = Field(..., description="Unique identifier for the employee")
    first_name: str = Field(..., description="First name of the employee")
    last_name: str = Field(..., description="Last name of the employee")
    title: str = Field(..., description="Job title of the employee")
    department: str = Field(..., description="Department of the employee")
    salary: float = Field(..., description="Annual salary of the employee")
    salary_band: str = Field(..., description="Salary band of the employee")
    hire_date: str = Field(..., description="Hire date of the employee")
    performance_review: str = Field(..., description="Annual performance review of the employee")
    payroll_data: dict = Field(..., description="Payroll data for the employee")
    employment_status: str = Field(..., description="Employment status of the employee")
    work_location: str = Field(..., description="Work location of the employee")
    remote: bool = Field(..., description="Whether the employee works remotely")
    benefits_enrollment: dict = Field(..., description="Benefits enrollment information")
    manager_id: str = Field(None, description="Employee ID of the manager")
    manager_name: str = Field(None, description="Name of the manager")
    email_address: str = Field(..., description="Email address of the employee")

async def generate_performance_review(first_name, last_name, title, company_data):
    prompt = f"""
You are a manager at {company_data['company_name']} preparing an annual performance review for {first_name} {last_name}, who is a {title}. Provide a realistic and professional performance review, highlighting strengths, areas for improvement, and goals for the next year. Avoid using any personal data or disallowed content.
"""
    try:
        # Use asyncio.to_thread to run the synchronous function asynchronously
        review = await asyncio.to_thread(chat_with_instructor, prompt, str)
        return review.strip() if review else "Performance review not available."
    except Exception as e:
        logging.error(f"Error generating performance review for {first_name} {last_name}: {e}")
        return "Performance review not available."

def generate_payroll_data(salary):
    # Generate payroll data similar to Gusto
    payroll = {
        'monthly_salary': round(salary / 12, 2),
        'taxes_withheld': round(salary * 0.2 / 12, 2),
        'benefits_deduction': round(salary * 0.05 / 12, 2),
        'net_pay': round((salary / 12) - (salary * 0.2 / 12) - (salary * 0.05 / 12), 2),
        'pay_period': 'Monthly',
        'last_pay_date': fake.date_between(start_date='-1m', end_date='today').isoformat()
    }
    return payroll

async def generate_employee(employee_info, company_data):
    first_name = employee_info['first_name']
    last_name = employee_info['last_name']
    title = employee_info['title']
    performance_review = await generate_performance_review(first_name, last_name, title, company_data)
    employee_info['performance_review'] = performance_review
    return employee_info

def generate_hr_data(company_data, employee_count=200):
    employees = []
    departments = ['Sales', 'Marketing', 'Engineering', 'HR', 'Finance', 'Support', 'Executive', 'Administration', 'Customer Experience']
    titles = {
        'Sales': ['Account Executive', 'Business Development Representative', 'Sales Manager', 'Sales Associate'],
        'Marketing': ['Marketing Manager', 'Content Strategist', 'SEO Specialist'],
        'Engineering': ['Software Engineer I', 'Software Engineer II', 'QA Engineer', 'Product Manager', 'UX/UI Designer'],
        'HR': ['HR Manager', 'Recruiter'],
        'Finance': ['Financial Analyst', 'Accountant', 'Finance Manager'],
        'Support': ['Support Specialist', 'Support Manager'],
        'Executive': ['Chief Executive Officer', 'Vice President of Sales', 'Vice President of Marketing', 'Vice President of Support', 'Vice President of Customer Experience'],
        'Administration': ['Office Manager'],
        'Customer Experience': ['Customer Success Manager']
    }

    # Define manager mapping
    manager_mapping = {}

    # CEO
    ceo = {
        'employee_id': 'EMP' + str(uuid.uuid4())[:8],
        'first_name': fake.first_name_female(),
        'last_name': fake.last_name(),
        'title': 'Chief Executive Officer',
        'department': 'Executive',
        'salary': round(random.uniform(200000, 300000), 2),
        'salary_band': 'Senior',
        'hire_date': fake.date_between(start_date='-10y', end_date='-5y').isoformat(),
        'employment_status': 'Full-Time',
        'work_location': 'San Francisco, CA',
        'remote': False,
        'benefits_enrollment': {
            'health_insurance': True,
            'dental_insurance': True,
            'retirement_plan': True
        },
        'manager_id': 'AC000',
        'manager_name': 'Board of Directors',
        'email_address': f"{'emily.smith'}@advancedcloud.com"
    }
    ceo['payroll_data'] = generate_payroll_data(ceo['salary'])
    employees.append(ceo)
    manager_mapping['CEO'] = ceo

    # Direct reports to CEO
    direct_reports_titles = [
        ('Vice President of Sales', 'Sales'),
        ('Vice President of Marketing', 'Marketing'),
        ('Vice President of Support', 'Support'),
        ('Vice President of Customer Experience', 'Customer Experience'),
        ('Finance Manager', 'Finance'),
        ('HR Manager', 'HR'),
        ('Engineering Manager', 'Engineering'),
        ('Office Manager', 'Administration')
    ]

    for title, department in direct_reports_titles:
        first_name = fake.first_name()
        last_name = fake.last_name()
        employee_id = 'EMP' + str(uuid.uuid4())[:8]
        manager = ceo
        salary = round(random.uniform(150000, 200000), 2)
        email_username = f"{first_name.lower()}.{last_name.lower()}"

        employee = {
            'employee_id': employee_id,
            'first_name': first_name,
            'last_name': last_name,
            'title': title,
            'department': department,
            'salary': salary,
            'salary_band': 'Senior',
            'hire_date': fake.date_between(start_date='-10y', end_date='-3y').isoformat(),
            'employment_status': 'Full-Time',
            'work_location': fake.city() + ', ' + fake.state_abbr(),
            'remote': random.choice([True, False]),
            'benefits_enrollment': {
                'health_insurance': True,
                'dental_insurance': True,
                'retirement_plan': True
            },
            'manager_id': manager['employee_id'],
            'manager_name': f"{manager['first_name']} {manager['last_name']}",
            'email_address': f"{email_username}@advancedcloud.com"
        }
        employee['payroll_data'] = generate_payroll_data(employee['salary'])
        employees.append(employee)
        manager_mapping[title] = employee

    # Generate other employees
    remaining_employee_count = employee_count - len(employees)
    for _ in range(remaining_employee_count):
        department = random.choice([dept for dept in departments if dept not in ['Executive', 'Administration']])
        title = random.choice(titles[department])
        first_name = fake.first_name()
        last_name = fake.last_name()
        employee_id = 'EMP' + str(uuid.uuid4())[:8]
        manager = manager_mapping.get(f"{department} Manager") or manager_mapping.get(f"Vice President of {department}") or ceo
        salary = round(random.uniform(50000, 120000), 2)
        salary_band = 'Junior' if 'I' in title else 'Mid' if 'II' in title else 'Senior'
        email_username = f"{first_name.lower()}.{last_name.lower()}"

        employee = {
            'employee_id': employee_id,
            'first_name': first_name,
            'last_name': last_name,
            'title': title,
            'department': department,
            'salary': salary,
            'salary_band': salary_band,
            'hire_date': fake.date_between(start_date='-5y', end_date='today').isoformat(),
            'employment_status': 'Full-Time',
            'work_location': fake.city() + ', ' + fake.state_abbr(),
            'remote': random.choice([True, False]),
            'benefits_enrollment': {
                'health_insurance': random.choice([True, False]),
                'dental_insurance': random.choice([True, False]),
                'retirement_plan': random.choice([True, False])
            },
            'manager_id': manager['employee_id'],
            'manager_name': f"{manager['first_name']} {manager['last_name']}",
            'email_address': f"{email_username}@advancedcloud.com"
        }
        employee['payroll_data'] = generate_payroll_data(employee['salary'])
        employees.append(employee)

    # Generate performance reviews asynchronously
    async def main():
        tasks = [generate_employee(emp, company_data) for emp in employees]
        return await asyncio.gather(*tasks)

    # Run the event loop
    employees = asyncio.run(main())

    # Save to CSV
    df_employees = pd.DataFrame(employees)
    df_employees.to_csv(os.path.join('data', 'employees.csv'), index=False)

    # Generate DOCX file
    generate_hr_records_docx(df_employees)

    logging.info("HR data generated.")
    print("HR data generated.")
    return employees

def generate_hr_records_docx(df_employees):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    document.add_heading('HR Records', 0)

    for idx, row in df_employees.iterrows():
        # Employee Name and ID as Heading
        employee_name = f"{row['first_name']} {row['last_name']} ({row['employee_id']})"
        document.add_heading(employee_name, level=1)

        # Employee Details
        details = document.add_paragraph()
        details.paragraph_format.space_after = Pt(6)

        # Fields to display
        fields_to_display = [
            ('Title', 'title'),
            ('Department', 'department'),
            ('Manager', 'manager_name'),
            ('Email', 'email_address'),
            ('Work Location', 'work_location'),
            ('Remote', 'remote'),
            ('Salary', 'salary'),
            ('Salary Band', 'salary_band'),
            ('Hire Date', 'hire_date'),
            ('Employment Status', 'employment_status'),
            ('Benefits Enrollment', 'benefits_enrollment'),
            ('Payroll Data', 'payroll_data'),
            ('Performance Review', 'performance_review')
        ]

        for field_name, field_key in fields_to_display:
            value = row.get(field_key, '')
            if isinstance(value, dict):
                # Format dictionaries nicely
                dict_items = '\n'.join([f"    {k}: {v}" for k, v in value.items()])
                text = f"{field_name}:\n{dict_items}"
            else:
                text = f"{field_name}: {value}"
            p = document.add_paragraph(text)
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_after = Pt(3)

    document.save(os.path.join('data', 'hr_records.docx'))
    print("HR records generated in 'data/hr_records.docx'.")
