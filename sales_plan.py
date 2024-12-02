# sales_plan.py

import os
import pandas as pd
from pydantic import BaseModel, Field
from typing import List
from chatgpt import chat_with_instructor
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import logging
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='sales_plan.log')

class SalesTarget(BaseModel):
    sales_rep_id: str = Field(..., description="Employee ID of the sales representative.")
    sales_rep_name: str = Field(..., description="Full name of the sales representative.")
    role: str = Field(..., description="Role of the sales representative (e.g., Account Executive, Sales Associate).")
    location: str = Field(..., description="Work location of the sales representative.")
    reports_to: str = Field(..., description="Manager's name.")
    total_sales_target: float = Field(..., description="Total sales target assigned to the representative.")
    new_sales_target: float = Field(..., description="Target for new sales (new clients).")
    upsell_target: float = Field(..., description="Target for upsells (existing clients).")
    strategies: List[str] = Field(..., description="Strategies for achieving the targets.")
    assigned_sales_associates: List[str] = Field(..., description="List of assigned Sales Associates (if any).")

class SalesPlan(BaseModel):
    company_name: str = Field(..., description="Name of the company.")
    total_revenue_target: float = Field(..., description="Total revenue target for the year.")
    new_sales_target: float = Field(..., description="Revenue target from new clients.")
    upsell_target: float = Field(..., description="Revenue target from existing clients.")
    sales_team_structure: List[dict] = Field(..., description="Structure of the sales team.")
    sales_targets: List[SalesTarget] = Field(..., description="Sales targets for each team member.")
    vp_responsibilities: List[str] = Field(..., description="Responsibilities of the VP of Sales.")
    action_plan: List[str] = Field(..., description="Strategies and action plans.")
    expected_outcomes: List[str] = Field(..., description="Expected outcomes from the plan.")
    contingency_plans: List[str] = Field(..., description="Contingency plans.")

async def generate_sales_plan_async(company_data):
    # Load data
    df_employees = pd.read_csv(os.path.join('data', 'employees.csv'))
    df_targets = pd.read_csv(os.path.join('data', 'sales_targets.csv'))

    # Prepare sales plan data
    sales_team_structure = []
    sales_targets = []
    total_revenue_target = 10000000
    new_sales_target = 6000000
    upsell_target = 4000000

    # Build sales team structure and sales targets
    for _, emp in df_employees.iterrows():
        if emp['department'] == 'Sales':
            role = emp['title']
            manager_name = emp['manager_name']
            sales_team_structure.append({
                'name': f"{emp['first_name']} {emp['last_name']} ({emp['employee_id']})",
                'position': role,
                'location': emp['work_location'],
                'reports_to': manager_name
            })

            # Get sales targets if available
            target_row = df_targets[df_targets['sales_rep_id'] == emp['employee_id']]
            if not target_row.empty:
                target_row = target_row.iloc[0]
                sales_target = SalesTarget(
                    sales_rep_id=emp['employee_id'],
                    sales_rep_name=f"{emp['first_name']} {emp['last_name']}",
                    role=role,
                    location=emp['work_location'],
                    reports_to=manager_name,
                    total_sales_target=target_row['total_sales_target'],
                    new_sales_target=target_row['new_clients_target'],
                    upsell_target=target_row['existing_clients_target'],
                    strategies=["Custom strategies will be included here."],
                    assigned_sales_associates=[]
                )
                sales_targets.append(sales_target)

    # Create sales plan object
    sales_plan_data = SalesPlan(
        company_name=company_data['company_name'],
        total_revenue_target=total_revenue_target,
        new_sales_target=new_sales_target,
        upsell_target=upsell_target,
        sales_team_structure=sales_team_structure,
        sales_targets=sales_targets,
        vp_responsibilities=[
            "Overall Accountability: Ensure the sales team meets the $10 million revenue target.",
            "Strategic Planning: Adjust strategies as needed.",
            "Team Support: Conduct performance reviews and foster collaboration."
        ],
        action_plan=[
            "Enhanced Collaboration",
            "Training and Development",
            "Market Segmentation",
            "Performance Incentives",
            "Resource Support"
        ],
        expected_outcomes=[
            "Revenue Achievement",
            "Increased Productivity",
            "Professional Growth",
            "Market Penetration"
        ],
        contingency_plans=[
            "Underperformance: Implement coaching plans.",
            "Resource Shifts: Reallocate resources as needed.",
            "Market Changes: Adjust strategies based on feedback."
        ]
    )

    # Prepare the prompt
    prompt = f"""
You are the Vice President of Sales at {sales_plan_data.company_name}. Based on the following sales plan data, create a comprehensive sales plan matching the format of the provided example.

Sales Plan Data:
{sales_plan_data.model_dump_json()}

Ensure the sales plan includes:
- Introduction
- Sales Team Structure
- Sales Targets
- Calculating Individual Targets
- Individual Sales Plans
- VP of Sales Responsibilities
- Action Plan and Strategies
- Expected Outcomes
- Contingency Plans
- Conclusion

Here is an example of the format:
Advanced Cloud 2024 Sales Plan (Revised)
________________


Introduction
The 2024 sales plan for Advanced Cloud aims to achieve a total of $10 million in new revenue. This target is divided into:
* $6 million from new sales (acquiring new clients).
* $4 million from upselling to existing clients (accounts existing as of January 1, 2024).
The sales team consists of:
* 6 Account Executives (AEs)
* 5 Sales Associates (SAs)
In this revised plan, each Sales Associate is assigned individual sales targets, which are 50% of the targets assigned to Account Executives. This plan cascades the overall targets down to every member of the sales organization, ensuring collective responsibility for achieving the company's goals.
________________


Sales Team Structure
Leadership
* Michael Johnson (AC002)
   * Position: Vice President of Sales
   * Location: New York, NY
   * Reports To: Emily Smith (CEO)
* Robert Miller (AC006)
   * Position: Sales Operations Manager
   * Location: Chicago, IL
   * Reports To: Michael Johnson (VP of Sales)
Account Executives
1. Richard Anderson (AC014)
   * Location: New York, NY
   * Reports To: Michael Johnson
2. Joseph Taylor (AC016)
   * Location: Boston, MA
   * Reports To: Michael Johnson
3. Charles Harris (AC020)
   * Location: Chicago, IL
   * Reports To: Michael Johnson
4. Michelle Lee (AC029)
   * Location: Austin, TX
   * Reports To: Michael Johnson
5. Jennifer King (AC035)
   * Location: Boston, MA
   * Reports To: Michael Johnson
6. Laura Phillips (AC049)
   * Location: Chicago, IL
   * Reports To: Michael Johnson
Sales Associates
1. Mary Thomas (AC015)
   * Location: Chicago, IL
   * Reports To: Assigned Account Executive(s)
2. Steven Lewis (AC028)
   * Location: New York, NY
   * Reports To: Assigned Account Executive(s)
3. Donald Hernandez (AC034)
   * Location: Chicago, IL
   * Reports To: Assigned Account Executive(s)
4. George Wright (AC036)
   * Location: New York, NY
   * Reports To: Assigned Account Executive(s)
5. Jason Campbell (AC050)
   * Location: Boston, MA
   * Reports To: Assigned Account Executive(s)
________________


Sales Targets
Overall Sales Targets
* Total Revenue Target: $10,000,000
   * New Sales (New Clients): $6,000,000
   * Upsells (Existing Clients): $4,000,000
Calculating Individual Targets
Total Target Units
* Account Executives (AEs): 6 AEs × 1 unit = 6 units
* Sales Associates (SAs): 5 SAs × 0.5 unit = 2.5 units
* Total Units: 6 + 2.5 = 8.5 units
Target per Unit
* Total Target: $10,000,000
* Target per Unit: $10,000,000 ÷ 8.5 units ≈ $1,176,470.59
Individual Targets
Account Executives
* Total Target per AE: 1 unit × $1,176,470.59 ≈ $1,176,471
   * New Sales Target (60%): $705,882
   * Upsell Target (40%): $470,588
Sales Associates
* Total Target per SA: 0.5 unit × $1,176,470.59 ≈ $588,235
   * New Sales Target (60%): $352,941
   * Upsell Target (40%): $235,294
Note: Targets are rounded to the nearest whole number for simplicity.
________________


Individual Sales Plans
Account Executives
1. Richard Anderson (AC014)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Focus on financial services and tech startups in New York.
   * Leverage existing network for referrals.
   * Partner with Sales Associate Steven Lewis.
2. Joseph Taylor (AC016)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Target biotech firms and universities in Boston.
   * Attend industry conferences.
   * Collaborate with Sales Associate Jason Campbell.
3. Charles Harris (AC020)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Focus on manufacturing and logistics companies in Chicago.
   * Implement customer loyalty programs.
   * Work with Sales Associate Mary Thomas.
4. Michelle Lee (AC029)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Engage with tech startups and SMEs in Austin.
   * Host monthly webinars.
   * Utilize local networking events.
5. Jennifer King (AC035)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Target healthcare and education sectors in Boston.
   * Develop strategic partnerships.
   * Supported by Sales Associate Jason Campbell.
6. Laura Phillips (AC049)
* Total Sales Target: $1,176,471
   * New Sales Target: $705,882
   * Upsell Target: $470,588
* Strategies:
   * Focus on retail and hospitality industries in Chicago.
   * Implement joint marketing initiatives.
   * Collaborate with Sales Associate Donald Hernandez.
Sales Associates
1. Mary Thomas (AC015)
* Total Sales Target: $588,235
   * New Sales Target: $352,941
   * Upsell Target: $235,294
* Assigned to: Charles Harris (AC020)
* Strategies:
   * Assist in lead generation and qualification.
   * Manage smaller accounts independently.
   * Focus on inside sales efforts.
2. Steven Lewis (AC028)
* Total Sales Target: $588,235
   * New Sales Target: $352,941
   * Upsell Target: $235,294
* Assigned to: Richard Anderson (AC014)
* Strategies:
   * Support outreach in the tech sector.
   * Handle initial product demos.
   * Nurture leads through the sales funnel.
3. Donald Hernandez (AC034)
* Total Sales Target: $588,235
   * New Sales Target: $352,941
   * Upsell Target: $235,294
* Assigned to: Laura Phillips (AC049)
* Strategies:
   * Focus on expanding into new markets.
   * Coordinate with marketing for lead campaigns.
   * Manage client onboarding processes.
4. George Wright (AC036)
* Total Sales Target: $588,235
   * New Sales Target: $352,941
   * Upsell Target: $235,294
* Assigned to: Michael Johnson (supporting multiple AEs)
* Strategies:
   * Provide sales analytics and reporting.
   * Assist in large account management.
   * Develop sales collateral.
5. Jason Campbell (AC050)
* Total Sales Target: $588,235
   * New Sales Target: $352,941
   * Upsell Target: $235,294
* Assigned to: Joseph Taylor (AC016) and Jennifer King (AC035)
* Strategies:
   * Manage client communications.
   * Schedule and coordinate sales meetings.
   * Contribute to proposal development.
________________


Revised Quarterly Targets
Account Executives
* Quarterly Total Target: $1,176,471 ÷ 4 ≈ $294,118 per quarter
   * Quarterly New Sales: $705,882 ÷ 4 ≈ $176,471
   * Quarterly Upsells: $470,588 ÷ 4 ≈ $117,647
Sales Associates
* Quarterly Total Target: $588,235 ÷ 4 ≈ $147,059 per quarter
   * Quarterly New Sales: $352,941 ÷ 4 ≈ $88,235
   * Quarterly Upsells: $235,294 ÷ 4 ≈ $58,824
________________


VP of Sales Responsibilities
Michael Johnson (AC002)
* Overall Accountability: Ensure the sales team meets the $10 million revenue target.
* Strategic Planning: Adjust strategies to incorporate Sales Associates' individual targets.
* Team Support:
   * Conduct bi-weekly performance reviews.
   * Foster collaboration between AEs and SAs.
   * Provide additional resources for training and development.
________________


Sales Operations Manager Role
Robert Miller (AC006)
* Data Analysis:
   * Monitor individual and team performance metrics.
   * Adjust forecasting models to include SAs' contributions.
* Process Optimization:
   * Enhance CRM workflows to track SA activities.
   * Implement new sales tools to support remote team members.
* Resource Allocation:
   * Coordinate with HR for any additional hiring needs.
   * Ensure equitable distribution of leads.
________________


Action Plan and Strategies
1. Enhanced Collaboration
* AE and SA Partnerships:
   * Establish clear roles and responsibilities.
   * Set up regular check-ins between AEs and their assigned SAs.
2. Training and Development
* Skill Enhancement:
   * Offer sales training programs tailored for SAs moving into target-carrying roles.
   * Provide mentorship opportunities.
3. Market Segmentation
* Target Allocation:
   * Assign SAs to specific market segments or smaller accounts.
   * Allow SAs to independently manage certain deals under AE supervision.
4. Performance Incentives
* Commission Structures:
   * Align compensation plans to reflect the new targets.
   * Offer bonuses for overachievement.
5. Resource Support
* Marketing Collaboration:
   * Create targeted campaigns to generate leads for both AEs and SAs.
* Technology Tools:
   * Utilize sales enablement platforms to streamline workflows.
________________


Expected Outcomes
* Revenue Achievement: Meet or exceed the $10 million revenue target for 2024.
* Increased Productivity: With SAs carrying targets, overall sales capacity increases.
* Professional Growth: Develop SAs into future AEs through experience and responsibility.
* Market Penetration: Broaden reach into smaller markets and accounts via SAs.
________________


Contingency Plans
* Underperformance:
   * Implement coaching plans for team members not meeting targets.
* Resource Shifts:
   * Reallocate SAs to different AEs if necessary to optimize performance.
* Market Changes:
   * Adjust strategies based on market feedback and economic indicators.
________________


Conclusion
By assigning individual sales targets to Sales Associates at 50% of the Account Executives' targets, Advanced Cloud leverages the full potential of its sales team. This collaborative approach not only aims to achieve the ambitious $10 million revenue goal but also fosters professional development and prepares the team for future growth.
________________


Prepared by: Advanced Cloud Sales Leadership Team


Use headings, subheadings, bullet points, and numbering where appropriate. Output the sales plan as plain text, formatted appropriately.
"""

    try:
        sales_plan_text = await asyncio.to_thread(chat_with_instructor, prompt, str)
        if sales_plan_text:
            sales_plan_text = sales_plan_text.strip()
            # Save sales plan as text file
            with open(os.path.join('data', 'sales_plan.txt'), 'w', encoding='utf-8') as f:
                f.write(sales_plan_text)
            print("Sales plan generated in 'data/sales_plan.txt'.")
            logging.info("Sales plan generated.")

            # Generate DOCX and PDF files
            generate_sales_plan_docx(sales_plan_text)
            generate_sales_plan_pdf(sales_plan_text)
        else:
            print("Failed to generate sales plan.")
            logging.error("Failed to generate sales plan.")
    except Exception as e:
        logging.error(f"Error generating sales plan: {e}")
        print("Error generating sales plan.")

def generate_sales_plan(company_data):
    asyncio.run(generate_sales_plan_async(company_data))

def generate_sales_plan_docx(sales_plan_text):
    document = Document()
    styles = document.styles

    # Modify styles for better formatting
    heading_style = styles['Heading 1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)

    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)

    lines = sales_plan_text.split('\n')
    for line in lines:
        line = line.strip()
        if line == '':
            continue
        elif line.endswith("(Revised)"):
            document.add_heading(line, level=0)
        elif line in [
            "Introduction", "Sales Team Structure", "Sales Targets",
            "Calculating Individual Targets", "Individual Sales Plans",
            "VP of Sales Responsibilities", "Action Plan and Strategies",
            "Expected Outcomes", "Contingency Plans", "Conclusion"
        ]:
            document.add_heading(line, level=1)
        elif line.startswith('*') or line.startswith('•') or line.startswith('-'):
            p = document.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line[0].isdigit() and (line[1] == '.' or line[1] == ')'):
            p = document.add_paragraph(line, style='List Number')
        else:
            document.add_paragraph(line)
    document.save(os.path.join('data', 'sales_plan.docx'))
    print("Sales plan generated in 'data/sales_plan.docx'.")

def generate_sales_plan_pdf(sales_plan_text):
    doc = SimpleDocTemplate(
        os.path.join('data', 'sales_plan.pdf'),
        pagesize=LETTER
    )
    styles = getSampleStyleSheet()
    # Use unique style names to avoid conflicts
    if 'CustomBullet' not in styles:
        styles.add(ParagraphStyle(name='CustomBullet', parent=styles['Normal'], bulletIndent=0, leftIndent=20))
    if 'CustomNumber' not in styles:
        styles.add(ParagraphStyle(name='CustomNumber', parent=styles['Normal'], bulletIndent=0, leftIndent=20))
    Story = []
    lines = sales_plan_text.split('\n')
    for line in lines:
        line = line.strip()
        if line == '':
            Story.append(Spacer(1, 12))
        elif line.endswith("(Revised)"):
            Story.append(Paragraph(line, styles['Title']))
        elif line in [
            "Introduction", "Sales Team Structure", "Sales Targets",
            "Calculating Individual Targets", "Individual Sales Plans",
            "VP of Sales Responsibilities", "Action Plan and Strategies",
            "Expected Outcomes", "Contingency Plans", "Conclusion"
        ]:
            Story.append(Paragraph(line, styles['Heading1']))
        elif line.startswith('*') or line.startswith('•') or line.startswith('-'):
            Story.append(Paragraph(line[1:].strip(), styles['CustomBullet']))
        elif line[0].isdigit() and (line[1] == '.' or line[1] == ')'):
            Story.append(Paragraph(line, styles['CustomNumber']))
        else:
            Story.append(Paragraph(line, styles['Normal']))
    doc.build(Story)
    print("Sales plan generated in 'data/sales_plan.pdf'.")

if __name__ == "__main__":
    company_data = {
        'company_name': 'Advanced Cloud'
    }
    generate_sales_plan(company_data)