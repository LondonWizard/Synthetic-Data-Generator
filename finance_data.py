# finance_data.py

import os
import pandas as pd
from faker import Faker
import random
import uuid
from datetime import datetime, timedelta
from docx import Document
from fpdf import FPDF

fake = Faker()
Faker.seed(0)

def generate_finance_data(company_data, customers, international_clients_percentage):
    invoices = []
    transactions = []
    num_invoices = 200

    for customer in customers:
        # Determine if the customer is international based on the percentage
        is_international = random.uniform(0, 100) < international_clients_percentage
        for _ in range(random.randint(1, 5)):  # Each customer can have multiple invoices
            invoice_id = f"INV{str(uuid.uuid4())[:8].upper()}"
            customer_id = customer['customer_id']
            issue_date = fake.date_between(start_date='-1y', end_date='today')
            due_date = issue_date + timedelta(days=30)
            amount = round(random.uniform(1000, 50000), 2)
            if is_international:
                currency = random.choice(['EUR', 'GBP', 'JPY', 'AUD'])
            else:
                currency = 'USD'
            status = random.choices(
                ['Paid', 'Unpaid', 'Overdue'],
                weights=[70, 20, 10],  # 70% paid on time
                k=1
            )[0]

            if currency != 'USD':
                # Introduce possible outstanding amounts due to currency conversions
                outstanding_amount = amount * random.uniform(0.01, 0.05) if status != 'Paid' else 0
            else:
                outstanding_amount = 0

            invoice = {
                'invoice_id': invoice_id,
                'customer_id': customer_id,
                'issue_date': issue_date.isoformat(),
                'due_date': due_date.isoformat(),
                'amount': amount,
                'currency': currency,
                'status': status,
                'outstanding_amount': round(outstanding_amount, 2)
            }
            invoices.append(invoice)

            # Generate transaction if invoice is paid
            if status == 'Paid':
                payment_date = issue_date + timedelta(days=random.randint(0, 30))
                transaction = {
                    'transaction_id': f"TRX{str(uuid.uuid4())[:8].upper()}",
                    'invoice_id': invoice_id,
                    'payment_date': payment_date.isoformat(),
                    'amount': amount,
                    'currency': currency
                }
                transactions.append(transaction)

    # Generate accounts payable and financial reports
    generate_accounts_payable(company_data)
    generate_financial_reports_document()

    # Save invoices to CSV
    df_invoices = pd.DataFrame(invoices)
    df_invoices.to_csv(os.path.join('data', 'invoices.csv'), index=False)

    # Save transactions to CSV
    df_transactions = pd.DataFrame(transactions)
    df_transactions.to_csv(os.path.join('data', 'transactions.csv'), index=False)

    # Generate Financial Reports
    generate_financial_reports_document()

    print("Finance data generated.")
    return invoices, transactions

def generate_accounts_payable(company_data):
    payables = []
    num_payables = 100

    for _ in range(num_payables):
        vendor_name = fake.company()
        invoice_id = f"AP{str(uuid.uuid4())[:8].upper()}"
        issue_date = fake.date_between(start_date='-1y', end_date='today')
        due_date = issue_date + timedelta(days=30)
        amount = round(random.uniform(500, 20000), 2)
        currency = 'USD'
        status = random.choice(['Paid', 'Unpaid', 'Overdue'])

        payable = {
            'invoice_id': invoice_id,
            'vendor_name': vendor_name,
            'issue_date': issue_date.isoformat(),
            'due_date': due_date.isoformat(),
            'amount': amount,
            'currency': currency,
            'status': status
        }
        payables.append(payable)

    # Save to CSV
    df_payables = pd.DataFrame(payables)
    df_payables.to_csv(os.path.join('data', 'accounts_payable.csv'), index=False)

    print("Accounts payable data generated.")
    return payables

def generate_financial_reports_document():
    # Generate 13-week Rolling Forecast
    forecast = []
    current_date = datetime.now().date()
    for i in range(13):
        week_date = current_date + timedelta(weeks=i)
        revenue = round(random.uniform(100000, 200000), 2)
        expenses = round(random.uniform(50000, 150000), 2)
        net_income = revenue - expenses

        forecast_entry = {
            'week_of': week_date.isoformat(),
            'revenue': revenue,
            'expenses': expenses,
            'net_income': net_income
        }
        forecast.append(forecast_entry)

    df_forecast = pd.DataFrame(forecast)
    df_forecast.to_csv(os.path.join('data', 'rolling_forecast.csv'), index=False)

    # Generate MRR (Monthly Recurring Revenue) file
    mrr = []
    for month_offset in range(-12, 1):
        month_date = current_date + timedelta(days=month_offset * 30)
        recurring_revenue = round(random.uniform(400000, 600000), 2)

        mrr_entry = {
            'month': month_date.strftime('%Y-%m'),
            'recurring_revenue': recurring_revenue
        }
        mrr.append(mrr_entry)

    df_mrr = pd.DataFrame(mrr)
    df_mrr.to_csv(os.path.join('data', 'mrr.csv'), index=False)

    # Prepare Financial Reports Content
    reports_content = f"""
Financial Reports

13-Week Rolling Forecast
-------------------------
"""

    for entry in forecast:
        reports_content += f"Week of {entry['week_of']}: Revenue: ${entry['revenue']:,.2f}, Expenses: ${entry['expenses']:,.2f}, Net Income: ${entry['net_income']:,.2f}\n"

    reports_content += "\nMRR (Monthly Recurring Revenue)\n-------------------------------\n"
    for entry in mrr:
        reports_content += f"{entry['month']}: ${entry['recurring_revenue']:,.2f}\n"

    # Generate PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', size=16)
    pdf.cell(200, 10, txt="Financial Reports", ln=True, align='C')
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=reports_content)
    pdf.output(os.path.join('data', 'financial_reports.pdf'))

    # Generate DOCX
    doc = Document()
    doc.add_heading('Financial Reports', 0)
    doc.add_heading('13-Week Rolling Forecast', level=1)
    for entry in forecast:
        doc.add_paragraph(f"Week of {entry['week_of']}: Revenue: ${entry['revenue']:,.2f}, Expenses: ${entry['expenses']:,.2f}, Net Income: ${entry['net_income']:,.2f}", style='List Bullet')
    doc.add_heading('MRR (Monthly Recurring Revenue)', level=1)
    for entry in mrr:
        doc.add_paragraph(f"{entry['month']}: ${entry['recurring_revenue']:,.2f}", style='List Bullet')
    doc.save(os.path.join('data', 'financial_reports.docx'))

    # Save Reports as TXT
    with open(os.path.join('data', 'financial_reports.txt'), 'w') as f:
        f.write(reports_content)

    print("Financial reports generated in PDF, DOCX, and TXT formats.")