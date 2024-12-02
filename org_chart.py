# org_chart.py

import os
import pandas as pd
from chatgpt import chat_with_instructor
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import LETTER
import logging
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='org_chart.log')

async def generate_org_chart_async(company_data):
    # Load employee data
    df_employees = pd.read_csv(os.path.join('data', 'employees.csv'))
    simplified_employees = df_employees[['employee_id', 'first_name', 'last_name', 'title', 'department', 'manager_name', 'work_location', 'remote']]
    employees = simplified_employees.to_dict('records')

    # Prepare the prompt
    prompt = f"""
You are an HR specialist at {company_data['company_name']}. Based on the following employee data, create an organizational chart in the specified format.

Employee Data:
{employees}

Format:
[Provide the desired format here, similar to the sample organizational chart provided earlier.]

Ensure that the chart includes:
- Hierarchical structure
- Employee names and IDs
- Positions and departments
- Locations and remote status
- Reporting lines

Here is an example of the expected format:

Advanced Cloud Organizational Chart
________________


1. Executive Leadership
Emily Smith (AC001)
* Position: Chief Executive Officer (CEO)
* Department: Executive
* Location: San Francisco, CA
* Remote: No
* Reports To: Board of Directors (AC000)
________________


2. Direct Reports to the CEO
A. James Garcia (AC008)
* Position: Engineering Manager
* Department: R&D/Product
* Location: San Francisco, CA
* Remote: No
* Reports To: Emily Smith (AC001)
B. Michael Johnson (AC002)
* Position: Vice President of Sales
* Department: Sales
* Location: New York, NY
* Remote: No
* Reports To: Emily Smith (AC001)
C. Linda Williams (AC003)
* Position: Vice President of Support
* Department: Support
* Location: Seattle, WA
* Remote: Yes
* Reports To: Emily Smith (AC001)
D. David Brown (AC004)
* Position: Vice President of Customer Experience
* Department: Customer Experience
* Location: Austin, TX
* Remote: No
* Reports To: Emily Smith (AC001)
E. Susan Jones (AC005)
* Position: Vice President of Marketing
* Department: Marketing
* Location: Boston, MA
* Remote: No
* Reports To: Emily Smith (AC001)
F. Donna Scott (AC039)
* Position: Finance Manager
* Department: Finance
* Location: Chicago, IL
* Remote: No
* Reports To: Emily Smith (AC001)
G. Michelle Adams (AC041)
* Position: HR Manager
* Department: Human Resources
* Location: Boston, MA
* Remote: No
* Reports To: Emily Smith (AC001)
H. Carol Nelson (AC043)
* Position: Office Manager
* Department: Administration
* Location: San Francisco, CA
* Remote: No
* Reports To: Emily Smith (AC001)
________________


3. Departmental Structure
A. R&D/Product Department
Led by James Garcia (AC008), Engineering Manager
Team Members:
1. Patricia Martinez (AC009)
   * Position: Software Engineer II
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
2. John Lopez (AC010)
   * Position: Software Engineer I
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: James Garcia (AC008)
3. Barbara Hernandez (AC011)
   * Position: Product Manager
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
4. William Gonzalez (AC012)
   * Position: QA Engineer
   * Location: Seattle, WA
   * Remote: Yes
   * Reports To: James Garcia (AC008)
5. Lisa Wilson (AC013)
   * Position: UX/UI Designer
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
6. Dorothy Martin (AC021)
   * Position: Software Engineer II
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
7. Christopher Thompson (AC022)
   * Position: QA Engineer
   * Location: Seattle, WA
   * Remote: Yes
   * Reports To: James Garcia (AC008)
8. Jessica Garcia (AC023)
   * Position: UX/UI Designer
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
9. Paul Walker (AC030)
   * Position: Software Engineer I
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
10. Laura Hall (AC031)
   * Position: Product Manager
   * Location: Seattle, WA
   * Remote: Yes
   * Reports To: James Garcia (AC008)
11. Mark Allen (AC032)
   * Position: Software Engineer II
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
12. Elizabeth Young (AC033)
   * Position: QA Engineer
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: James Garcia (AC008)
13. Deborah Mitchell (AC045)
   * Position: Software Engineer I
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
14. Ronald Perez (AC046)
   * Position: Software Engineer II
   * Location: Seattle, WA
   * Remote: Yes
   * Reports To: James Garcia (AC008)
15. Sharon Roberts (AC047)
   * Position: UX/UI Designer
   * Location: San Francisco, CA
   * Remote: No
   * Reports To: James Garcia (AC008)
16. Anthony Turner (AC048)
   * Position: QA Engineer
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: James Garcia (AC008)
________________


B. Sales Department
Led by Michael Johnson (AC002), VP of Sales
Direct Report:
* Robert Miller (AC006)
   * Position: Sales Operations Manager
   * Location: Chicago, IL
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
Sales Team:
1. Richard Anderson (AC014)
   * Position: Account Executive
   * Location: New York, NY
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
2. Mary Thomas (AC015)
   * Position: Sales Associate
   * Location: Chicago, IL
   * Remote: No
   * Reports To: Michael Johnson (AC002)
3. Joseph Taylor (AC016)
   * Position: Account Executive
   * Location: Boston, MA
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
4. Charles Harris (AC020)
   * Position: Account Executive
   * Location: Chicago, IL
   * Remote: No
   * Reports To: Michael Johnson (AC002)
5. Steven Lewis (AC028)
   * Position: Sales Associate
   * Location: New York, NY
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
6. Michelle Lee (AC029)
   * Position: Account Executive
   * Location: Austin, TX
   * Remote: No
   * Reports To: Michael Johnson (AC002)
7. Donald Hernandez (AC034)
   * Position: Sales Associate
   * Location: Chicago, IL
   * Remote: No
   * Reports To: Michael Johnson (AC002)
8. Jennifer King (AC035)
   * Position: Account Executive
   * Location: Boston, MA
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
9. George Wright (AC036)
   * Position: Sales Associate
   * Location: New York, NY
   * Remote: No
   * Reports To: Michael Johnson (AC002)
10. Laura Phillips (AC049)
   * Position: Account Executive
   * Location: Chicago, IL
   * Remote: No
   * Reports To: Michael Johnson (AC002)
11. Jason Campbell (AC050)
   * Position: Sales Associate
   * Location: Boston, MA
   * Remote: Yes
   * Reports To: Michael Johnson (AC002)
________________


C. Support Department
Led by Linda Williams (AC003), VP of Support
Team Members:
1. Daniel Martinez (AC024)
   * Position: Support Engineer
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: Linda Williams (AC003)
2. Ashley Robinson (AC025)
   * Position: Support Engineer
   * Location: Seattle, WA
   * Remote: No
   * Reports To: Linda Williams (AC003)
________________


D. Customer Experience Department
Led by David Brown (AC004), VP of Customer Experience
Team Members:
1. Matthew Clark (AC026)
   * Position: Customer Success Manager
   * Location: Chicago, IL
   * Remote: Yes
   * Reports To: David Brown (AC004)
2. Amy Rodriguez (AC027)
   * Position: Customer Success Associate
   * Location: Boston, MA
   * Remote: No
   * Reports To: David Brown (AC004)
________________


E. Marketing Department
Led by Susan Jones (AC005), VP of Marketing
Direct Report:
* Karen Davis (AC007)
   * Position: Marketing Operations Manager
   * Location: New York, NY
   * Remote: Yes
   * Reports To: Susan Jones (AC005)
Marketing Team:
1. Sarah Moore (AC017)
   * Position: Marketing Specialist
   * Location: New York, NY
   * Remote: No
   * Reports To: Susan Jones (AC005)
2. Thomas Jackson (AC018)
   * Position: Content Strategist
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: Susan Jones (AC005)
3. Margaret White (AC019)
   * Position: SEO Analyst
   * Location: Boston, MA
   * Remote: No
   * Reports To: Susan Jones (AC005)
4. Sandra Lopez (AC037)
   * Position: Marketing Specialist
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: Susan Jones (AC005)
5. Kenneth Hill (AC038)
   * Position: Content Strategist
   * Location: Boston, MA
   * Remote: No
   * Reports To: Susan Jones (AC005)
________________


F. Finance Department
Led by Donna Scott (AC039), Finance Manager
Team Members:
1. Edward Green (AC040)
   * Position: Accountant
   * Location: Seattle, WA
   * Remote: Yes
   * Reports To: Donna Scott (AC039)
________________


G. Human Resources Department
Led by Michelle Adams (AC041), HR Manager
Team Members:
1. Kevin Baker (AC042)
   * Position: Recruiter
   * Location: New York, NY
   * Remote: Yes
   * Reports To: Michelle Adams (AC041)
________________


H. Administration Department
Led by Carol Nelson (AC043), Office Manager
Team Members:
1. Brian Carter (AC044)
   * Position: Administrative Assistant
   * Location: Austin, TX
   * Remote: Yes
   * Reports To: Carol Nelson (AC043)

Output the full organizational chart as plain text. Do not trim the list of any employee, and provide the full org chart outline. Include all employees that are inputted in this prompt in the output
"""

    try:
        org_chart_text = await asyncio.to_thread(chat_with_instructor, prompt, str)
        if org_chart_text:
            org_chart_text = org_chart_text.strip()
            # Save org chart as text file
            with open(os.path.join('data', 'org_chart.txt'), 'w', encoding='utf-8') as f:
                f.write(org_chart_text)
            print("Organizational chart generated in 'data/org_chart.txt'.")
            logging.info("Organizational chart generated.")

            # Generate DOCX and PDF files
            generate_org_chart_docx(org_chart_text)
            generate_org_chart_pdf(org_chart_text)
        else:
            print("Failed to generate organizational chart.")
            logging.error("Failed to generate organizational chart.")
    except Exception as e:
        logging.error(f"Error generating organizational chart: {e}")
        print("Error generating organizational chart.")

def generate_org_chart(company_data):
    asyncio.run(generate_org_chart_async(company_data))

def generate_org_chart_docx(org_chart_text):
    document = Document()
    for line in org_chart_text.split('\n'):
        if line.strip() == '':
            document.add_paragraph('')
        elif line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.'):
            document.add_heading(line.strip(), level=1)
        elif line.strip()[0].isalpha() and line.strip()[1] == '.':
            document.add_heading(line.strip(), level=2)
        elif line.strip()[0].isdigit() and '.' in line:
            p = document.add_paragraph(line.strip())
            p.paragraph_format.left_indent = Pt(36)
        elif line.strip().startswith('*') or line.strip().startswith('-'):
            p = document.add_paragraph(line.strip())
            p.paragraph_format.left_indent = Pt(72)
        else:
            document.add_paragraph(line.strip())
    document.save(os.path.join('data', 'org_chart.docx'))
    print("Organizational chart generated in 'data/org_chart.docx'.")

def generate_org_chart_pdf(org_chart_text):
    doc = SimpleDocTemplate(
        os.path.join('data', 'org_chart.pdf'),
        pagesize=LETTER
    )
    styles = getSampleStyleSheet()
    # Use unique style names to avoid conflicts
    if 'CustomBullet' not in styles:
        styles.add(ParagraphStyle(name='CustomBullet', parent=styles['Normal'], bulletIndent=0, leftIndent=20))
    Story = []
    lines = org_chart_text.split('\n')
    for line in lines:
        line = line.strip()
        if line == '':
            Story.append(Spacer(1, 12))
        elif line.startswith('•') or line.startswith('*') or line.startswith('-'):
            Story.append(Paragraph(line[1:].strip(), styles['CustomBullet']))
        else:
            Story.append(Paragraph(line, styles['Normal']))
    doc.build(Story)
    print("Organizational chart generated in 'data/org_chart.pdf'.")

if __name__ == "__main__":
    company_data = {
        'company_name': 'Advanced Cloud'
    }
    generate_org_chart(company_data)