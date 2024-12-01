# sales_data.py

import os
import pandas as pd
from pydantic import BaseModel, Field
from faker import Faker
import random
import uuid
import asyncio
import logging
from datetime import datetime, timedelta
from chatgpt import chat_with_instructor
from docx import Document
from fpdf import FPDF

fake = Faker()
Faker.seed(0)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='sales_data.log')

class Opportunity(BaseModel):
    opportunity_id: str = Field(..., description="Unique identifier for the sales opportunity")
    sales_rep_id: str = Field(..., description="ID of the sales representative")
    client_name: str = Field(..., description="Name of the client")
    stage: str = Field(..., description="Current stage of the sales process")
    amount: float = Field(..., description="Potential deal amount")
    close_date: str = Field(..., description="Expected close date")
    renewal_date: str = Field(..., description="Expected renewal date")
    email_conversations: str = Field(..., description="Email conversation thread with the client")

async def generate_email_conversation(sales_rep_name, contact_name, client_name, stage, company_name):
    prompt = f"""
You are a sales representative named {sales_rep_name} at {company_name}. Generate a professional email conversation between you and a client named {contact_name} from {client_name}. The conversation should revolve around advancing a sales opportunity currently at the '{stage}' stage, discussing details relevant to this stage. Avoid using any personal data or disallowed content. Provide the email thread in chronological order, starting with the client's email and including your responses. Each email should have a date in the format 'YYYY-MM-DD'.
"""
    try:
        email_conversation = await asyncio.to_thread(chat_with_instructor, prompt, str)
        return email_conversation.strip() if email_conversation else None
    except Exception as e:
        logging.error(f"Error generating email conversation for {client_name}: {e}")
        return None

async def generate_opportunity(sales_rep, company_data, stages):
    sales_rep_id = sales_rep['employee_id']
    sales_rep_name = f"{sales_rep['first_name']} {sales_rep['last_name']}"
    client_name = fake.company()
    stage = random.choice(stages)
    amount = round(random.uniform(50000, 150000), 2)
    close_date_obj = fake.date_between(start_date='-1y', end_date='today')
    close_date = close_date_obj.isoformat()

    if isinstance(close_date_obj, datetime):
        close_date_obj = close_date_obj.date()

    renewal_date_obj = close_date_obj + timedelta(days=365)
    renewal_date = renewal_date_obj.isoformat()

    contact_name = fake.name()
    email_conversation = await generate_email_conversation(
        sales_rep_name, contact_name, client_name, stage, company_name=company_data['company_name']
    )
    if email_conversation is None:
        return None

    opportunity = Opportunity(
        opportunity_id=f"OPP{str(uuid.uuid4())}",
        sales_rep_id=sales_rep_id,
        client_name=client_name,
        stage=stage,
        amount=amount,
        close_date=close_date,
        renewal_date=renewal_date,
        email_conversations=email_conversation
    )
    return opportunity

def generate_customers_from_opportunities(opportunities):
    customers = []
    seen_clients = set()
    for opp in opportunities:
        client_name = opp.client_name
        if client_name not in seen_clients:
            customer = {
                'customer_id': f"CUST{str(uuid.uuid4())[:8].upper()}",
                'company_name': client_name,
                # Add other customer fields as needed
            }
            customers.append(customer)
            seen_clients.add(client_name)
    return customers

def generate_sales_data(company_data, employees, num_opportunities=200):
    opportunities = []
    stages = ['Prospecting', 'Qualification', 'Proposal', 'Negotiation', 'Closed Won', 'Closed Lost']

    # Filter for sales representatives
    sales_reps = [emp for emp in employees if emp['department'] == 'Sales']

    if not sales_reps:
        logging.error("No Sales representatives available to generate sales data.")
        print("Error: No Sales representatives available to generate sales data.")
        return [], [], []

    async def main():
        tasks = [
            generate_opportunity(random.choice(sales_reps), company_data, stages)
            for _ in range(num_opportunities)
        ]
        return await asyncio.gather(*tasks)

    # Use asyncio.run() instead of get_event_loop()
    results = asyncio.run(main())

    # Filter out None results
    opportunities = [opp for opp in results if opp is not None]

    # Generate customers from opportunities
    customers = generate_customers_from_opportunities(opportunities)

    # Define sales_team
    sales_team = sales_reps

    # Save customers to CSV
    df_customers = pd.DataFrame(customers)
    df_customers.to_csv(os.path.join('data', 'customers.csv'), index=False)

    # Save opportunities to CSV
    df_opportunities = pd.DataFrame([opp.dict() for opp in opportunities])
    df_opportunities.to_csv(os.path.join('data', 'opportunities.csv'), index=False)

    # Generate Sales Opportunity Report
    generate_sales_opportunity_report(opportunities, sales_team, company_data)

    logging.info("Sales data generated.")
    print("Sales data generated.")
    return customers, sales_team, opportunities

def generate_sales_opportunity_report(opportunities, sales_team, company_data):
    # Prepare Report Content
    report_content = f"""
Advanced Cloud Sales Opportunity Report

Total Opportunities: {len(opportunities)}
Stages Distribution:
"""

    stage_counts = {}
    for opp in opportunities:
        stage_counts[opp.stage] = stage_counts.get(opp.stage, 0) + 1

    for stage, count in stage_counts.items():
        # Replace '•' with '*' to avoid UnicodeEncodeError
        report_content += f"* {stage}: {count}\n"

    # Generate PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', size=16)
    pdf.cell(200, 10, txt="Advanced Cloud Sales Opportunity Report", ln=True, align='C')
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=report_content)
    pdf.output(os.path.join('data', 'sales_opportunity_report.pdf'))

    # Generate DOCX
    doc = Document()
    doc.add_heading('Advanced Cloud Sales Opportunity Report', 0)
    doc.add_paragraph(f"Total Opportunities: {len(opportunities)}")
    doc.add_heading('Stages Distribution:', level=1)
    for stage, count in stage_counts.items():
        doc.add_paragraph(f"{stage}: {count}", style='List Bullet')
    doc.save(os.path.join('data', 'sales_opportunity_report.docx'))

    # Save Report as TXT
    with open(os.path.join('data', 'sales_opportunity_report.txt'), 'w') as f:
        f.write(report_content)

    print("Sales Opportunity Report generated in PDF, DOCX, and TXT formats.")