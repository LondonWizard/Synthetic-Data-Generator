# ae_handbook.py

import os
from pydantic import BaseModel, Field
from typing import List
from chatgpt import chat_with_instructor
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import logging
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', filename='ae_handbook.log')

class AEHandbook(BaseModel):
    company_name: str = Field(..., description="The name of the company.")
    last_updated: str = Field(..., description="The last updated date of the handbook.")
    sections: List[str] = Field(..., description="List of section titles to include in the handbook.")

async def generate_ae_handbook_async(company_data):
    # Prepare the handbook data
    last_updated = '11.24.2023'
    sections = [
        'Introduction',
        'Company Overview',
        'Daily Workflow Overview',
        'Managing Leads',
        'Deal Management',
        'Opportunity Management',
        'Follow-Up Communication',
        'Setting and Achieving Close Targets',
        'Time Management and Productivity',
        'Utilizing Sales Tools and CRM Systems',
        'Best Practices and Professional Development',
        'Frequently Asked Questions',
        'Appendices',
        'Resources',
        'Contact Information'
    ]

    ae_handbook_data = AEHandbook(
        company_name=company_data['company_name'],
        last_updated=last_updated,
        sections=sections
    )

    # Prepare the prompt
    prompt = f"""
You are tasked with creating an Account Executive Handbook for {ae_handbook_data.company_name}. The handbook should be comprehensive and match the following format:

Advanced Cloud
Account Executive Handbook
Last updated: {ae_handbook_data.last_updated}

[Section Titles and Content]

Ensure that each section provides detailed guidance tailored to {ae_handbook_data.company_name}'s policies and procedures. Use headings, subheadings, bullet points, and numbering where appropriate to enhance readability.

Sections to include:
{ae_handbook_data.sections}


Here is an example of the expected format:

Advanced Cloud 
Account Executive Handbook 
Last updated: 11.24.2023
Introduction
Welcome to Advanced Cloud, where we revolutionize how businesses operate through cutting-edge cloud solutions. As an Account Executive, you play a pivotal role in driving our company's growth by connecting businesses with our SaaS offerings.
This handbook is your comprehensive guide to succeeding in your role. It covers daily workflows, best practices, and detailed instructions on managing every aspect of the sales process.
________________


Company Overview
Advanced Cloud is a leading SaaS provider specializing in scalable cloud solutions for businesses of all sizes. Our mission is to empower organizations by simplifying their operations and enhancing productivity through innovative technology.
________________


Daily Workflow Overview
A structured day enhances productivity and ensures that you meet your targets. Below is a suggested daily schedule:
1. Morning Preparation (8:00 AM - 8:30 AM)
   * Review emails and voicemails.
   * Check your calendar for scheduled meetings.
2. Team Sync-Up (8:30 AM - 9:00 AM)
   * Attend daily stand-up meetings.
   * Share updates and align on priorities.
3. Lead Management (9:00 AM - 11:00 AM)
   * Review new leads assigned.
   * Prioritize follow-ups based on lead scoring.
4. Client Meetings and Demos (11:00 AM - 1:00 PM)
   * Conduct scheduled meetings.
   * Provide product demonstrations.
5. Lunch Break (1:00 PM - 2:00 PM)
6. Deal Progression (2:00 PM - 3:30 PM)
   * Update deal statuses in CRM.
   * Prepare proposals and contracts.
7. Opportunity Management (3:30 PM - 4:30 PM)
   * Identify upselling opportunities.
   * Engage with existing clients.
8. Follow-Up Communication (4:30 PM - 5:30 PM)
   * Send follow-up emails.
   * Make necessary phone calls.
9. End-of-Day Review (5:30 PM - 6:00 PM)
   * Plan for the next day.
   * Log daily activities.
________________


Managing Leads
Effective lead management is crucial for converting prospects into clients.
4.1 Lead Generation
* Inbound Leads:
   * Monitor leads coming from marketing campaigns.
   * Respond promptly to inquiries.
* Outbound Prospecting:
   * Use tools like LinkedIn Sales Navigator.
   * Attend networking events.
4.2 Lead Qualification
* BANT Framework:
   * Budget: Assess if the prospect has the financial capacity.
   * Authority: Confirm decision-makers.
   * Need: Identify specific pain points.
   * Timeline: Understand their purchase timeframe.
* Discovery Calls:
   * Prepare questions in advance.
   * Document key insights.
4.3 Lead Prioritization
* Lead Scoring:
   * Use CRM to assign scores based on engagement.
   * Focus on high-scoring leads first.
* Segmentation:
   * Categorize leads by industry, size, or potential value.
   * Tailor your approach accordingly.
4.4 Lead Nurturing
* Content Sharing:
   * Send relevant case studies.
   * Provide informational webinars.
* Regular Check-Ins:
   * Schedule follow-up calls.
   * Keep the conversation ongoing.
________________


Deal Management
Managing deals effectively moves prospects through the sales funnel.
5.1 Creating a New Deal
* CRM Entry:
   * Fill in all mandatory fields: deal name, contact info, estimated value.
   * Assign a probability percentage.
* Associations:
   * Link the deal to relevant contacts and companies.
   * Attach previous interaction notes.
5.2 Deal Stages and Progression
* Deal Stages:
   * Prospecting
   * Qualification
   * Proposal
   * Negotiation
   * Closing
* Advancing Deals:
   * Update the stage after each significant interaction.
   * Set reminders for next steps.
5.3 Deal Negotiation Strategies
* Preparation:
   * Know your bottom line.
   * Understand the client's priorities.
* Communication:
   * Be clear and concise.
   * Use data to support your position.
* Concessions:
   * Offer value-added services instead of price reductions.
   * Get something in return for concessions made.
5.4 Risk Management
* Identifying Red Flags:
   * Non-responsive contacts.
   * Unreasonable demands.
* Mitigation Strategies:
   * Address concerns proactively.
   * Involve senior team members when necessary.
________________


Opportunity Management
Opportunities are potential deals that can significantly impact revenue.
6.1 Identifying Opportunities
* Market Analysis:
   * Stay informed about industry trends.
   * Identify gaps our solutions can fill.
* Client Feedback:
   * Listen to existing clients for new needs.
   * Use satisfaction surveys.
6.2 Opening Opportunities
* Opportunity Creation:
   * Log in CRM with all relevant details.
   * Assign to the appropriate pipeline.
* Initial Engagement:
   * Reach out to key stakeholders.
   * Schedule introductory meetings.
6.3 Opportunity Tracking
* KPIs:
   * Monitor metrics like conversion rates.
   * Adjust strategies based on data.
* Reporting:
   * Provide regular updates to management.
   * Use dashboards for visualization.
6.4 Closing Opportunities
* Final Presentations:
   * Customize proposals to client needs.
   * Highlight ROI and value proposition.
* Contract Negotiation:
   * Work with legal and finance teams.
   * Ensure terms are mutually beneficial.
* Onboarding Handoff:
   * Introduce the client to the customer success team.
   * Provide all necessary documentation.
________________


Follow-Up Communication
Timely and effective communication keeps prospects engaged.
7.1 Effective Email Follow-Ups
* Subject Lines:
   * Keep them concise and relevant.
   * Example: "Next Steps After Our Meeting"
* Email Body:
   * Personalize the message.
   * Include a clear call-to-action.
* Timing:
   * Send follow-ups within 24 hours.
   * Schedule subsequent emails appropriately.
7.2 Phone Call Best Practices
* Preparation:
   * Review previous interactions.
   * Set a clear objective for the call.
* During the Call:
   * Use active listening.
   * Address any objections.
* After the Call:
   * Summarize key points in an email.
   * Outline agreed-upon next steps.
7.3 Meeting Recaps and Next Steps
* Recap Emails:
   * Send immediately after meetings.
   * Highlight decisions and action items.
* Scheduling:
   * Propose dates for the next meeting.
   * Use calendar invites to confirm.
________________


Setting and Achieving Close Targets
Meeting your sales quotas is essential for personal and company success.
8.1 Understanding Your Quotas
* Monthly and Quarterly Goals:
   * Know your targets for each period.
   * Break them down into weekly objectives.
* Product Mix:
   * Be aware of which products or services to focus on.
   * Align with company priorities.
8.2 Pipeline Management
* Pipeline Health:
   * Regularly assess the stages of your deals.
   * Ensure a steady flow of prospects.
* Forecasting:
   * Use historical data to predict outcomes.
   * Adjust your focus based on likelihood to close.
8.3 Performance Tracking
* Use of CRM Analytics:
   * Monitor your activities and outcomes.
   * Identify areas for improvement.
* Feedback Loops:
   * Seek input from managers.
   * Implement suggested strategies.
8.4 Strategies for Meeting Targets
* Time Allocation:
   * Focus on high-value activities.
   * Limit time spent on low-probability deals.
* Skill Development:
   * Attend training sessions.
   * Practice sales pitches.
* Collaboration:
   * Work with marketing for lead generation.
   * Engage with product teams for technical insights.
________________


Time Management and Productivity
Efficient time management maximizes your effectiveness.
* Prioritization:
   * Use the Eisenhower Matrix to categorize tasks.
   * Focus on important and urgent activities.
* Automation Tools:
   * Utilize email templates.
   * Set up automated reminders.
* Avoiding Distractions:
   * Limit time on non-work-related activities.
   * Use productivity apps to stay focused.
________________


Utilizing Sales Tools and CRM Systems
Our tools are designed to support your sales efforts.
* CRM Usage:
   * Log all interactions.
   * Update deal stages promptly.
* Sales Enablement Platforms:
   * Access the latest sales materials.
   * Share content with prospects.
* Analytics Tools:
   * Review dashboards for insights.
   * Monitor key performance indicators.
________________


Best Practices and Professional Development
Continuous improvement is key to success.
* Networking:
   * Build relationships within the industry.
   * Attend conferences and webinars.
* Learning Opportunities:
   * Complete mandatory training modules.
   * Pursue certifications relevant to your role.
* Mentorship:
   * Seek guidance from experienced colleagues.
   * Participate in mentorship programs.
________________


Frequently Asked Questions
1. How quickly should I follow up after a meeting?
   * Within 24 hours to maintain engagement.
2. What if a client raises an objection I can't address?
   * Note the concern and consult with your manager or the relevant department.
3. How do I update my close targets?
   * Discuss adjustments during your one-on-one meetings with your sales manager.
________________


Appendices
A. Email Templates
* Introduction Email
* Meeting Request
* Follow-Up After No Response
B. Call Scripts
* Discovery Call Script
* Objection Handling Script
* Closing Call Script
C. CRM Workflow Guide
* How to Create a Lead
* Updating Deal Stages
* Generating Reports
D. Sales Metrics and KPIs
* Definitions of Key Metrics
* Monthly Reporting Templates
* Benchmarking Data
________________


Resources
* Sales Team SharePoint
* Product Knowledge Base
* Competitor Analysis Reports
* Marketing Collateral Library
________________


Contact Information
* Sales Manager: [Name], [Email], [Phone]
* HR Department: [Email], [Phone]
* IT Support: [Email], [Phone]
________________


Empower Businesses with Advanced Cloud Solutions
Your role as an Account Executive is vital to our mission. By following this guide and utilizing the resources provided, you'll be well-equipped to achieve your targets and contribute to the success of both our clients and Advanced Cloud.
________________


This handbook is intended for internal use only. Please refer to the company intranet for the most up-to-date information.

Output the handbook as plain text, formatted appropriately.
"""

    try:
        handbook_text = await asyncio.to_thread(chat_with_instructor, prompt, str)
        if handbook_text:
            handbook_text = handbook_text.strip()
            # Save as DOCX and PDF
            generate_handbook_docx(handbook_text, ae_handbook_data)
            generate_handbook_pdf(handbook_text, ae_handbook_data)
            print("Account Executive Handbook generated.")
            logging.info("Account Executive Handbook generated.")
        else:
            print("Failed to generate Account Executive Handbook.")
            logging.error("Failed to generate Account Executive Handbook.")
    except Exception as e:
        logging.error(f"Error generating Account Executive Handbook: {e}")
        print("Error generating Account Executive Handbook.")

def generate_ae_handbook(company_data):
    asyncio.run(generate_ae_handbook_async(company_data))

def generate_handbook_docx(handbook_text, ae_handbook_data):
    document = Document()
    styles = document.styles

    # Modify styles for better formatting
    heading_style = styles['Heading 1']
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)

    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)

    lines = handbook_text.split('\n')
    for line in lines:
        line = line.strip()
        if line == '':
            continue
        elif line == ae_handbook_data.company_name or line == 'Account Executive Handbook':
            continue  # Already added
        elif line.startswith('Last updated:'):
            continue  # Already added
        elif line in ae_handbook_data.sections:
            document.add_heading(line, level=2)
        elif line.startswith('•') or line.startswith('*') or line.startswith('-'):
            p = document.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line[0].isdigit() and (line[1] == '.' or line[1] == ')'):
            p = document.add_paragraph(line, style='List Number')
        else:
            document.add_paragraph(line)
    document.save(os.path.join('data', 'ae_handbook.docx'))
    print("Account Executive Handbook generated in 'data/ae_handbook.docx'.")

def generate_handbook_pdf(handbook_text, ae_handbook_data):
    doc = SimpleDocTemplate(
        os.path.join('data', 'ae_handbook.pdf'),
        pagesize=LETTER
    )
    styles = getSampleStyleSheet()
    # Use unique style names to avoid conflicts
    if 'CustomBullet' not in styles:
        styles.add(ParagraphStyle(name='CustomBullet', parent=styles['Normal'], bulletIndent=0, leftIndent=20))
    if 'CustomNumber' not in styles:
        styles.add(ParagraphStyle(name='CustomNumber', parent=styles['Normal'], bulletIndent=0, leftIndent=20))
    Story = []
    Story.append(Paragraph(ae_handbook_data.company_name, styles['Title']))
    Story.append(Paragraph('Account Executive Handbook', styles['Title']))
    Story.append(Paragraph(f"Last updated: {ae_handbook_data.last_updated}", styles['Normal']))
    Story.append(Spacer(1, 12))

    lines = handbook_text.split('\n')
    for line in lines:
        line = line.strip()
        if line == '':
            Story.append(Spacer(1, 12))
        elif line == ae_handbook_data.company_name or line == 'Account Executive Handbook':
            continue  # Already added
        elif line.startswith('Last updated:'):
            continue  # Already added
        elif line in ae_handbook_data.sections:
            Story.append(Paragraph(line, styles['Heading1']))
        elif line.startswith('•') or line.startswith('*') or line.startswith('-'):
            Story.append(Paragraph(line[1:].strip(), styles['CustomBullet']))
        elif line[0].isdigit() and (line[1] == '.' or line[1] == ')'):
            Story.append(Paragraph(line, styles['CustomNumber']))
        else:
            Story.append(Paragraph(line, styles['Normal']))
    doc.build(Story)
    print("Account Executive Handbook generated in 'data/ae_handbook.pdf'.")

if __name__ == "__main__":
    company_data = {
        'company_name': 'Advanced Cloud'
    }
    generate_ae_handbook(company_data)